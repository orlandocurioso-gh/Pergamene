def creaCamicia(pergamena):
    import Moduli
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    nome=pergamena.nomeCognome
    pathCamiciaMaster='C:\\Users\\orlan\\Documents\\Coding\\Python\\Sapienza\\Pergamene\\masters\\'
    pathDiSalvataggio='C:\\Users\\orlan\\Documents\\Coding\\Python\\Sapienza\\Pergamene\\camicie\\'
    pathFirme='C:\\Users\\orlan\\Documents\\Coding\\Python\\Sapienza\\Pergamene\\firme\\'
    camiciaMaster='camiciaMasterTesto.docx'
    d=Document(pathCamiciaMaster+camiciaMaster)
    for p in d.paragraphs:
        if '<<' in p.text:
            match p.text:
                case '<<facolta>>':
                    testoApp=pergamena.facolta
                    p.clear()
                    runner=p.add_run(testoApp)
                    runner.bold=True
                    runner.font.name='Times New Roman'
                    runner.font.size=Pt(14)
                case '<<classeLaurea>>':
                    testoApp=pergamena.classeLaurea
                    p.clear()
                    runner=p.add_run(testoApp)
                    runner.bold=True
                    runner.font.name='Times New Roman'
                case '<<nomeCognome>>':
                    testoApp='A      '+pergamena.nomeCognome
                    p.clear()
                    runner=p.add_run(testoApp)
                    runner.font.name='Times New Roman'
                    runner.font.size=Pt(12)
                case '<<sesso>> <<luogoNascita>> <<dataNascita>>':
                    if pergamena.sesso=='m':
                        nascita='nato'
                    else:
                        nascita='nata'
                    testoApp='         '+nascita+' a '+pergamena.luogoNascita+' '+pergamena.provincia+' il '+pergamena.dataNascita
                    p.clear()
                    runner=p.add_run(testoApp)
                    runner.font.name='Times New Roman'
                    runner.font.size=Pt(12)
                case '<<firmaRettrice>>':
                    p.clear()
                    p.add_run().add_picture(pathFirme+pergamena.firmaRettrice+'.tif', width=Pt(150))
                case '<<firmaDg>><<firmaPreside>>':
                    p.clear()
                    p.add_run().add_picture(pathFirme+pergamena.firmaDg+'.tif', width=Pt(150))
                    p.add_run().add_text('			                       ')
                    p.add_run().add_picture(pathFirme+pergamena.firmaPreside+'.tif', width=Pt(150))
                case '<<dataLaurea>>':
                    p.text=p.text.replace('<<dataLaurea>>','                                        '+pergamena.dataLaurea)
                case '<<protocollo>>':
                    p.text=p.text.replace('<<protocollo>>','N. DI PROTOCOLLO      '+pergamena.protocollo)
                case '<<dataConsegna>>':
                    p.text=p.text.replace('<<dataConsegna>>','DATO IN ROMA IL        '+pergamena.dataConsegna)
                case '<<numeroDiploma>>':
                    p.text=p.text.replace('<<numeroDiploma>>','Diploma n.:'+pergamena.numeroDiploma)
            #print (p.text)
    d.save(pathDiSalvataggio+'camicia'+nome+'.docx')